import os
import glob
import json
from docx import Document
import EnumerablesFiles
from py_linq import Enumerable
import xml.etree.ElementTree as ET

names = []
with open('Name.json') as f:
    names = json.load(f)

values = names[1]
items = names[2]
names = names[0]

"""
def findProcentResult(path):
    result = []
    for order in EnumerablesFiles.EnumerableOrder(path):
        print("order: " + order)
        result.append(0)
        for docx in EnumerablesFiles.EnumerableDoc(path + order):
            result[-1]+=1
    print("empity result = " + str(Enumerable(result).where(lambda count: count == 0).count()))
    print("count = " +  str(Enumerable(result).count()))
    return
"""

def GetElement(arr, root):
    element = ET.SubElement(root, "node")
    if isinstance(arr, str):
        element.text = arr
        return element
    if isinstance(arr, float):
        element.text = arr
        return element
    for node in arr:
        GetElement(node, element)
    return element

def CreateXml(path, arr):
    # we make root element
    root = ET.Element("node")
    GetElement(arr, root)
    tree = ET.ElementTree(root)

    # write the tree into an XML file
    tree.write(path, encoding='utf-8', xml_declaration=True)

def Main(path):
    for order in EnumerablesFiles.EnumerableOrder(path):
        print("order: " + order)
        arr = []
        for docx in EnumerablesFiles.EnumerableDocx(path + order):
            result, data = BuildData(docx)
            if result is True:
                arr.append(data)

                info = []
                for name in data:
                    info.append(isColumnType(name, names))
                arr.append(info)
        print(arr)
        CreateXml(path + order + ".xml", arr)
    return


def SaveJson(path, data):
    with open(path, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=4)
    return


def BuildData(path):
    print("Read ->" + path)
    try:
        Document(path)
    except Exception as e:
        print(e)
        return False, None
    result = []
    for t in Document(path).tables:
        tab = []
        try:
            header = []
            for i in range(0, t._column_count):
                if t.rows[0] is None or t._cells is None:
                    continue
                if parsed(t.rows[0].cells[i].text):
                    header.append(i)

            columns_count = len(header)
            if columns_count > 2:
                try:
                    for row in t.rows:
                        row_data = []
                        if row.cells[columns_count - 1].text:
                            for col in header:
                                row_data.append(row.cells[col].text)

                        if row_data:
                            tab.append(row_data)
                except Exception as e:
                    print(e)
        except Exception as e:
            print(e)

        if tab:
            result.append(tab)

    if result:
        return True, result
    else:
        return False, None


def parsed(str):
    info = []
    info.append(isColumnType(str, names))
    info.append(isColumnType(str, values))
    info.append(isColumnType(str, items))
    return Enumerable(info).max() > 0


def writeRow():
    return


def isColumnType(name, indificators):
    count = 0.0
    for indificator in indificators:
        if indificator in name:
            count += 1.0
    return count / len(indificators)  # , count/name.split()


Main("data/")
